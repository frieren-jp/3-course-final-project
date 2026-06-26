from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document


WORK = Path(".docx-audit")
FILES = {
    "PM02_source": WORK / "PM02_source.docx",
    "PM11_source": WORK / "PM11_source.docx",
    "PM11_org_sample": WORK / "PM11_org_sample.docx",
}
MAJOR = {"ЗАДАНИЕ", "ДНЕВНИК", "АТТЕСТАЦИОННЫЙ ЛИСТ", "ХАРАКТЕРИСТИКА", "ОТЧЕТ"}


def attrs(xml: str, tag: str) -> dict[str, str]:
    m = re.search(fr"<w:{tag}\b([^>]*)/?>(?:</w:{tag}>)?", xml)
    if not m:
        return {}
    return dict(re.findall(r"w:([A-Za-z0-9]+)=\"([^\"]*)\"", m.group(1)))


def break_before(p) -> bool:
    if "<w:pageBreakBefore" in p._p.xml:
        return True
    prev = p._p.getprevious()
    while prev is not None:
        xml = prev.xml
        text = "".join(prev.xpath(".//w:t/text()")).strip()
        if "<w:br" in xml and 'w:type="page"' in xml:
            return True
        if text:
            return False
        prev = prev.getprevious()
    return False


def run_marks(doc: Document) -> list[dict]:
    out = []
    for pi, p in enumerate(doc.paragraphs):
        for run in p.runs:
            f = run.font
            marks = []
            if f.highlight_color is not None:
                marks.append(f"highlight={f.highlight_color}")
            if run.bold:
                marks.append("bold")
            if run.italic:
                marks.append("italic")
            if run.underline:
                marks.append("underline")
            if f.color is not None and f.color.rgb is not None:
                marks.append(f"color={f.color.rgb}")
            if marks and run.text.strip():
                out.append({"p": pi, "text": run.text.strip()[:80], "marks": marks})
    return out


def audit(label: str, path: Path) -> dict:
    doc = Document(str(path))
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
        styles = z.read("word/styles.xml").decode("utf-8", errors="ignore") if "word/styles.xml" in z.namelist() else ""
    major = []
    ps = list(doc.paragraphs)
    for i, p in enumerate(ps):
        t = p.text.strip()
        if t in MAJOR:
            org_idx = None
            for j in range(i - 1, max(-1, i - 12), -1):
                if "Автономная некоммерческая" in ps[j].text:
                    org_idx = j
                    break
            major.append(
                {
                    "idx": i,
                    "text": t,
                    "title_break_before": break_before(p),
                    "block_start_idx": org_idx,
                    "block_break_before": break_before(ps[org_idx]) if org_idx is not None else None,
                }
            )
    return {
        "label": label,
        "path": str(path.resolve()),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "sectPr_count": xml.count("<w:sectPr"),
        "pgSz": attrs(xml, "pgSz"),
        "pgMar": attrs(xml, "pgMar"),
        "xml_counts": {
            "highlight": xml.count("<w:highlight"),
            "shading": xml.count("<w:shd"),
            "trHeight": xml.count("<w:trHeight"),
            "caps": xml.count("<w:caps"),
            "smallCaps": xml.count("<w:smallCaps"),
        },
        "styles_counts": {
            "highlight": styles.count("<w:highlight"),
            "shading": styles.count("<w:shd"),
        },
        "major": major,
        "direct_marks_sample": run_marks(doc)[:30],
    }


def main() -> None:
    data = [audit(label, path) for label, path in FILES.items()]
    out = WORK / "source_audit.json"
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(out.resolve())


if __name__ == "__main__":
    main()
