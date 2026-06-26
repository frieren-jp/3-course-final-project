from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


WORK = Path(".docx-audit")

MAJOR_TITLES = {
    "ЗАДАНИЕ",
    "ДНЕВНИК",
    "АТТЕСТАЦИОННЫЙ ЛИСТ",
    "ХАРАКТЕРИСТИКА",
    "ОТЧЕТ",
}

FORCED_NEW_PAGE_TITLES = {
    # The first block already naturally starts after the title page in Word.
    # Forcing it creates an empty page 2. These are the blocks that otherwise
    # tend to merge with the previous block in the filled reports.
    "ДНЕВНИК",
    "ХАРАКТЕРИСТИКА",
    "ОТЧЕТ",
}

SKIP_REPORT_LABELS = {
    "Дата сдачи отчета",
    "Подпись руководителя практики",
    "ФИО\t\t\tподпись",
    "(ФИО)                           (подпись)",
}


def set_plain_tnr12(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = False
    run.italic = False
    run.underline = False
    run.font.highlight_color = None
    if run.font.color is not None:
        run.font.color.rgb = None


def start_on_new_page(paragraph) -> None:
    # Use Word's paragraph property instead of an empty paragraph with a manual
    # page break; empty break paragraphs can create accidental blank pages.
    paragraph.paragraph_format.page_break_before = True


def ensure_major_blocks_start_new_page(doc: Document, extra_titles: set[str] | None = None) -> None:
    forced_titles = set(FORCED_NEW_PAGE_TITLES)
    if extra_titles:
        forced_titles.update(extra_titles)
    paragraphs = list(doc.paragraphs)
    targets = []
    for i, p in enumerate(paragraphs):
        if p.text.strip() not in forced_titles:
            continue
        block_start_idx = None
        for j in range(i - 1, max(-1, i - 12), -1):
            if "Автономная некоммерческая" in paragraphs[j].text:
                block_start_idx = j
                break
        if block_start_idx is not None and block_start_idx > 0:
            targets.append(block_start_idx)

    for idx in sorted(set(targets), reverse=True):
        start_on_new_page(paragraphs[idx])


def remove_trailing_empty_paragraphs(doc: Document) -> None:
    body = doc._body._element
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        body.remove(doc.paragraphs[-1]._element)


def normalize_report_body(doc: Document) -> None:
    report_idx = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "ОТЧЕТ"), None)
    if report_idx is None:
        return

    # Keep the official title layout, but remove accidental emphasis from the subtitle.
    if report_idx + 1 < len(doc.paragraphs):
        for run in doc.paragraphs[report_idx + 1].runs:
            set_plain_tnr12(run)

    for p in doc.paragraphs[report_idx + 2 :]:
        text = p.text.strip()
        if not text or text in SKIP_REPORT_LABELS:
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_plain_tnr12(run)


def patch_page_setup(path: Path) -> None:
    tmp = path.with_suffix(".tmp.docx")
    sect_pr = (
        '<w:sectPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:pgSz w:w="11906" w:h="16838" w:orient="portrait"/>'
        '<w:pgMar w:top="1134" w:right="567" w:bottom="1134" w:left="1701" '
        'w:header="720" w:footer="720" w:gutter="0"/>'
        "</w:sectPr>"
    )
    pg_mar = (
        '<w:pgMar w:top="1134" w:right="567" w:bottom="1134" w:left="1701" '
        'w:header="720" w:footer="720" w:gutter="0"/>'
    )
    pg_sz = '<w:pgSz w:w="11906" w:h="16838" w:orient="portrait"/>'

    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                text = data.decode("utf-8")
                if "<w:sectPr" not in text:
                    text = text.replace("</w:body>", sect_pr + "</w:body>")
                else:
                    if re.search(r"<w:pgSz[^>]*/>", text):
                        text = re.sub(r"<w:pgSz[^>]*/>", pg_sz, text, count=1)
                    else:
                        text = text.replace("<w:sectPr", pg_sz + "<w:sectPr", 1)
                    if re.search(r"<w:pgMar[^>]*/>", text):
                        text = re.sub(r"<w:pgMar[^>]*/>", pg_mar, text, count=1)
                    else:
                        text = text.replace("</w:sectPr>", pg_mar + "</w:sectPr>", 1)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def fix_one(src: str, dst: str) -> Path:
    doc = Document(str(WORK / src))
    extra_titles = {"АТТЕСТАЦИОННЫЙ ЛИСТ"} if src == "PM11_source.docx" else set()
    ensure_major_blocks_start_new_page(doc, extra_titles)
    normalize_report_body(doc)
    remove_trailing_empty_paragraphs(doc)
    out = WORK / dst
    doc.save(str(out))
    patch_page_setup(out)
    return out


if __name__ == "__main__":
    for output in (
        fix_one("PM02_source.docx", "PM02_checked.docx"),
        fix_one("PM11_source.docx", "PM11_checked.docx"),
    ):
        print(output.resolve())
